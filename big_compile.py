'''
Created on Mar 10, 2021

@author: austin
'''
import docx
from docx import Document
from functions.format_preserver import format_preserver2

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def big_compile(doc_names, output_name):
    the_big_one = Document()
    shells = Document()
    links = Document()
    impacts = Document()
    alts = Document()
    answersto = Document()
    for name in doc_names:
        doc = Document(name)
        delete_paragraph(doc.paragraphs[0])
        while len(doc.paragraphs) > 0:
            if doc.paragraphs[0].style.name == 'Heading 1':
                break
            else:
                format_preserver2(shells, doc.paragraphs[0])
                delete_paragraph(doc.paragraphs[0])
        delete_paragraph(doc.paragraphs[0])
        while len(doc.paragraphs) > 0:
            if doc.paragraphs[0].style.name == 'Heading 1':
                break
            else:
                format_preserver2(links, doc.paragraphs[0])
                delete_paragraph(doc.paragraphs[0])
        delete_paragraph(doc.paragraphs[0])
        while len(doc.paragraphs) > 0:
            if doc.paragraphs[0].style.name == 'Heading 1':
                break
            else:
                format_preserver2(impacts, doc.paragraphs[0])
                delete_paragraph(doc.paragraphs[0])
        delete_paragraph(doc.paragraphs[0])
        while len(doc.paragraphs) > 0:
            if doc.paragraphs[0].style.name == 'Heading 1':
                break
            else:
                format_preserver2(alts, doc.paragraphs[0])
                delete_paragraph(doc.paragraphs[0])
        delete_paragraph(doc.paragraphs[0])
        while len(doc.paragraphs) > 0:
            if doc.paragraphs[0].style.name == 'Heading 1':
                break
            else:
                format_preserver2(answersto, doc.paragraphs[0])
                delete_paragraph(doc.paragraphs[0])

    the_big_one.add_paragraph('Shells','Heading 1')
    for p in shells.paragraphs:
        format_preserver2(the_big_one, p)
    the_big_one.add_paragraph('Links','Heading 1')
    for p in links.paragraphs:
        format_preserver2(the_big_one, p)
    the_big_one.add_paragraph('Impacts','Heading 1')
    for p in impacts.paragraphs:
        format_preserver2(the_big_one, p)
    the_big_one.add_paragraph('Alts','Heading 1')
    for p in alts.paragraphs:
        format_preserver2(the_big_one, p)
    the_big_one.add_paragraph('Answers To','Heading 1')
    for p in answersto.paragraphs:
        format_preserver2(the_big_one, p)
    the_big_one.save(output_name)
    return

x = [r"C:\Users\austi\Documents\setcolproj\processed once\7.docx", r"C:\Users\austi\Documents\setcolproj\processed once\5.docx"]
y = r"C:\Users\austi\Documents\setcolproj\processed once\new test.docx"
big_compile(x, y)